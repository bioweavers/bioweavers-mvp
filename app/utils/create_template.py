from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_pto_template(output_path="../app/utils/pto_template.docx"):
    doc = Document()

    doc.add_heading("PTO Assessment Report", level=1)
    doc.add_heading("Potential to Occur Determinations", level=2)

    table = doc.add_table(rows=11, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = [
        "Scientific Name\nCommon Name",
        "Status",
        "Habitat Requirements",
        "Potential to Occur",
        "Habitat Suitability/\nObservations",
    ]

    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h

    # CNPS section header
    cnps_header = table.rows[1].cells
    cnps_header[0].merge(cnps_header[-1])
    cnps_header[0].text = "CNPS Species"

    # CNPS loop start
    table.rows[2].cells[0].text = "{%tr for row in cnps_rows %}"

    # CNPS data row
    cnps_data = table.rows[3].cells
    cnps_data[0].text = "{{ row.SpeciesDisplay }}"
    cnps_data[1].text = "{{ row.StatusDisplay }}"
    cnps_data[2].text = "{{ row.HabitatRequirements }}"
    cnps_data[3].text = "{{ row.PotentialtoOccur }}"
    cnps_data[4].text = "{{ row.HabitatSuitabilityObservations }}"

    # CNPS loop end
    table.rows[4].cells[0].text = "{%tr endfor %}"

    # Blank spacer row if you want
    spacer = table.rows[5].cells
    spacer[0].merge(spacer[-1])
    spacer[0].text = ""

    # CNDDB section header
    cnddb_header = table.rows[6].cells
    cnddb_header[0].merge(cnddb_header[-1])
    cnddb_header[0].text = "CNDDB Species"

    # CNDDB loop start
    table.rows[7].cells[0].text = "{%tr for row in cnddb_rows %}"

    # CNDDB data row
    cnddb_data = table.rows[8].cells
    cnddb_data[0].text = "{{ row.SpeciesDisplay }}"
    cnddb_data[1].text = "{{ row.StatusDisplay }}"
    cnddb_data[2].text = "{{ row.HabitatRequirements }}"
    cnddb_data[3].text = "{{ row.PotentialtoOccur }}"
    cnddb_data[4].text = "{{ row.HabitatSuitabilityObservations }}"

    # CNDDB loop end
    table.rows[9].cells[0].text = "{%tr endfor %}"

    # Optional ending blank row
    end = table.rows[10].cells
    end[0].merge(end[-1])
    end[0].text = ""

    doc.save(output_path)