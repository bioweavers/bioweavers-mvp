"""
Create a Word template with Jinja2 placeholders for docxtpl.

This script creates the template programmatically. In practice, you'd 
design this in Word and add the placeholders manually. This is just 
to bootstrap the example without requiring Word.

Run once: python create_template.py
"""

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

def create_pto_template():
    doc = Document()
    
    # Title
    title = doc.add_heading('PTO Assessment Report', level=1)
    
    # Project info paragraph with placeholders
    doc.add_paragraph('Project: {{ project_name }}')
    doc.add_paragraph('Location: {{ project_location }}')
    doc.add_paragraph('Buffer Radius: {{ buffer_radius }}')
    doc.add_paragraph('Assessment Date: {{ assessment_date }}')
    doc.add_paragraph()
    
    # Summary
    doc.add_heading('Species Assessment Summary', level=2)
    doc.add_paragraph('Total species assessed: {{ total_species }}')
    doc.add_paragraph('Species with potential to occur: {{ species_with_potential }}')
    doc.add_paragraph()
    
    # The PTO Table
    doc.add_heading('Potential to Occur Determinations', level=2)
    
    # Create table: header + 3 rows for the loop pattern
    # Row pattern for docxtpl table loops:
    #   Row 0: Headers (static)
    #   Row 1: {%tr for ... %} - loop start (gets removed)
    #   Row 2: Data template row (gets cloned N times)
    #   Row 3: {%tr endfor %} - loop end (gets removed)
    table = doc.add_table(rows=4, cols=5)
    table.style = 'Table Grid'
    
    # Row 0: Header row
    headers = ['Scientific Name', 'Common Name', 'Status', 'PTO', 'Rationale']
    hdr_cells = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr_cells[i].text = header
        for paragraph in hdr_cells[i].paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # Row 1: Loop start (this row gets removed after processing)
    loop_start = table.rows[1].cells
    loop_start[0].text = '{%tr for species in species_data %}'
    
    # Row 2: Data template row (this gets cloned for each item)
    data_row = table.rows[2].cells
    data_row[0].text = '{{ species.scientific_name }}'
    data_row[1].text = '{{ species.common_name }}'
    data_row[2].text = '{{ species.status }}'
    data_row[3].text = '{{ species.pto }}'
    data_row[4].text = '{{ species.rationale }}'
    
    # Row 3: Loop end (this row gets removed after processing)
    loop_end = table.rows[3].cells
    loop_end[0].text = '{%tr endfor %}'
    
    # Footer note
    doc.add_paragraph()
    doc.add_paragraph('PTO Definitions:', style='Heading 3')
    doc.add_paragraph('• Present: Species documented on site')
    doc.add_paragraph('• High: Suitable habitat present, known occurrences nearby')
    doc.add_paragraph('• Moderate: Some suitable habitat, occurrences in region')
    doc.add_paragraph('• Low: Marginal habitat, few regional occurrences')
    doc.add_paragraph('• Not Expected: Outside range or no suitable habitat')
    
    # Save
    doc.save('pto_template.docx')
    print("✓ Created pto_template.docx")

if __name__ == '__main__':
    create_pto_template()
