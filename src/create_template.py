from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT


def create_pto_template(output_path="src/pto_template.docx"):
    """
        Create the Word template used for the PTO Assessment Report export.

        This function builds a .docx template programmatically using python-docx.
        The template is designed to work with docxtpl, which later fills the table
        with grouped Potential to Occur data from the Streamlit app.

        The generated template includes:
        - A main report heading
        - A section heading
        - A table with fixed column headers
        - docxtpl loop tags for taxon groups
        - docxtpl loop tags for species rows within each taxon group

        Parameters
        ----------
        output_path : str, optional
            File path where the generated Word template should be saved.
            Default is "src/pto_template.docx".

        Returns
        -------
        None
            The function saves the Word template to the given output path.

        Template Context Requirements
        -----------------------------
        The template expects a context dictionary with this structure:

            {
                "taxon_groups": [
                    {
                        "TaxonGroup": "Birds",
                        "rows": [
                            {
                                "SpeciesDisplay": "...",
                                "StatusDisplay": "...",
                                "HabitatRequirements": "...",
                                "PotentialtoOccur": "...",
                                "HabitatSuitabilityObservations": "..."
                            }
                        ]
                    }
                ]
            }

        Notes
        -----
        This function only creates the blank template. It does not insert real PTO
        data. The actual data is inserted later by make_buffer() using docxtpl.
        """

    # Create a new blank Word document
    doc = Document()

    # Add report headings
    doc.add_heading("PTO Assessment Report", level=1)
    doc.add_heading("Potential to Occur Determinations", level=2)

    # Create the base table including:
        # 0: header row
        # 1: taxon group loop start
        # 2: taxon group heading
        # 3: species row loop start
        # 4: species data row
    table = doc.add_table(rows=5, cols=5)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Define column headers
    headers = [
        "Scientific Name\nCommon Name",
        "Status",
        "Habitat Requirements",
        "Potential to Occur",
        "Habitat Suitability/\nObservations",
    ]

    # Populate the header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h

    # Start the outer docxtpl loop
    # This repeats the taxon group section for each group in taxon_groups.
    table.rows[1].cells[0].text = "{%tr for group in taxon_groups %}"

    # Create the taxon group header row
    # The first cell is merged across the full table width so the taxon group name
    # appears as a section heading.
    group_header = table.rows[2].cells
    group_header[0].merge(group_header[-1])
    group_header[0].text = "{{ group.TaxonGroup }}"

    # Start the inner  docxtpl loop
    # This repeats one species row for each row inside the current taxon group
    table.rows[3].cells[0].text = "{%tr for row in group.rows %}"

    # Add placeholders for one species data row
    # Names must match keys produced by format_cnps(), format_cnddb(),
    # and make_buffer()
    data = table.rows[4].cells
    data[0].text = "{{ row.SpeciesDisplay }}"
    data[1].text = "{{ row.StatusDisplay }}"
    data[2].text = "{{ row.HabitatRequirements }}"
    data[3].text = "{{ row.PotentialtoOccur }}"
    data[4].text = "{{ row.HabitatSuitabilityObservations }}"

    # End the inner species row loop
    row_end = table.add_row().cells
    row_end[0].text = "{%tr endfor %}"

    # End the outer taxon group loop
    group_end = table.add_row().cells
    group_end[0].text = "{%tr endfor %}"

    # Save the generated template
    doc.save(output_path)


if __name__ == "__main__":
    # Allows this file to be run directly to regenerate the Word template
    create_pto_template()
